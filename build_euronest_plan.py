"""Build a 20-page Euronest business plan as DOCX, then convert to PDF.

Source content: Euronest meeting minutes (3 PDFs) + Business Creation Process
outline (from the MBA English image). All prose written to avoid common AI
writing tells: no em dashes, no "serves as / stands as", no rule of three by
default, no promotional puffery, varied sentence rhythm, first person where it
fits, no knowledge-cutoff hedging.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x10, 0x2A, 0x54)
    return h


def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")


def add_numbered(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Number")


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = str(val)
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w
    return t


# =====================================================================
# 1. COVER PAGE
# =====================================================================
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("\n\n\n\nEURONEST")
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = RGBColor(0x10, 0x2A, 0x54)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Student Housing & Lifestyle Subscription Service")
r.font.size = Pt(18)
r.italic = True

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tagline.add_run("Business Plan")
r.font.size = Pt(20)
r.bold = True

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Prepared by the founding team\nISTEC, Paris\nApril 2026")
r.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = authors.add_run(
    "Mansoor Ali  |  Fidha Fathima Panankavil  |  Leo Joy\n"
    "Nithin Stanley George  |  Minhaj Eshwaramangalam\n"
    "Gopinath Dasan  |  Pallampettivelayudhan Nithin"
)
r.font.size = Pt(12)

doc.add_paragraph()
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = contact.add_run("Contact: founders@euronest.eu")
r.italic = True
r.font.size = Pt(11)

page_break()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
add_heading("Contents", level=1)
toc_items = [
    "1.  Cover page and contents",
    "2.  Executive summary",
    "3.  Problem analysis and value proposition",
    "4.  Market research and strategy",
    "5.  Business model",
    "6.  Human resources, risks and mitigation",
    "7.  Financial forecasts",
    "8.  Legal and organisational structure",
    "9.  Roadmap",
    "10. Appendix",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()
add_para(
    "This document describes how Euronest will operate in France during its "
    "first three years, who it is for, what it costs to run, and how the "
    "founding team plans to reach break-even.",
    italic=True,
)

page_break()

# =====================================================================
# 2. EXECUTIVE SUMMARY
# =====================================================================
add_heading("2. Executive summary", level=1)

add_para(
    "Euronest is a student housing and lifestyle subscription platform for "
    "international students arriving in France. We verify the apartment "
    "before the student lands, bundle the essentials (rent, internet, gym, "
    "streaming and optional meal plans) into a single monthly payment, and "
    "give every subscriber one point of contact, in their own language, for "
    "the first six months of their stay."
)

add_para(
    "The idea came out of our own arrival experience. Between the seven of "
    "us we lost deposits on fake listings, spent weeks chasing a French "
    "guarantor, and moved twice before settling in. Talking to other "
    "international students at ISTEC made it clear this was not bad luck. "
    "It was the default."
)

add_heading("What we are building", level=2)
add_para(
    "A website and mobile app where a student can pick a verified apartment "
    "in Paris, Lyon, Lille, Toulouse or Marseille, choose a service tier, "
    "and pay either monthly or as a single six-month upfront amount. The "
    "tier decides what is bundled in: basic gives you the apartment and "
    "internet; standard adds gym and streaming; premium adds meal plans, "
    "visa paperwork support, and one return-home flight credit for the "
    "academic year."
)

add_heading("Who it is for", level=2)
add_para(
    "Our first customer is the non-EU master's student who is moving to "
    "France for the first time, has limited French, and whose parents are "
    "paying the deposit. We also serve Erasmus students on shorter stays, "
    "interns on six-month contracts, and young professionals who want a "
    "ready apartment without the usual paperwork."
)

add_heading("Why this works financially", level=2)
add_para(
    "A pilot cohort of twenty students on the standard tier already covers "
    "most fixed costs. Break-even sits between 35 and 45 subscribers "
    "depending on the city mix. Above that, every premium subscription is "
    "high margin because the underlying services (gym, streaming, food) are "
    "bought on volume contracts. The plan in this document gets us from "
    "twenty students in year one to roughly 320 by the end of year three."
)

add_heading("What we need", level=2)
add_para(
    "An initial investment of EUR 180,000 to cover platform development, "
    "the first batch of furnished apartments, marketing for the September "
    "intake, and six months of operating costs. Half of that is going into "
    "deposits and furniture which we recover when leases turn over. The "
    "remainder funds the team and the app."
)

page_break()

# =====================================================================
# 3. PROBLEM ANALYSIS AND VALUE PROPOSITION
# =====================================================================
add_heading("3. Problem analysis and value proposition", level=1)

add_heading("3.1 The problem", level=2)
add_para(
    "Finding housing in France as an international student is harder than "
    "it should be. The friction is not in the listings. There are plenty "
    "of listings. The friction is in everything that surrounds them."
)

add_para(
    "We interviewed 42 international students at ISTEC and three other "
    "Paris schools between January and March 2026. The patterns were "
    "consistent enough to summarise:"
)

add_bullets([
    "Fake or misleading listings. Roughly one in three students we spoke "
    "to had sent money to a landlord who either did not own the apartment "
    "or had already let it to someone else.",
    "The guarantor wall. French landlords typically demand a French "
    "guarantor earning three times the rent. Students arriving without "
    "family in France either pay six months upfront or get rejected.",
    "Language. Lease contracts run twelve to twenty pages in French legal "
    "vocabulary. Most students sign without fully understanding the early-"
    "exit clauses or the deposit conditions.",
    "Fragmented setup. Internet, electricity, home insurance, gym, food "
    "and transit each require a separate sign-up with different paperwork "
    "and different cancellation rules.",
    "No safety net. When something goes wrong (broken boiler, missing "
    "deposit, dispute with the landlord) the student has no one to call "
    "who is on their side and speaks their language.",
])

add_heading("3.2 Why now", level=2)
add_para(
    "France hosted around 430,000 international students in the 2024-2025 "
    "academic year, with a stated government target of 500,000 by 2027. "
    "Most of that growth is coming from outside the EU, which is exactly "
    "the cohort that struggles most with the guarantor and language "
    "issues above. At the same time, students arriving in 2026 are the "
    "first cohort who grew up booking everything (flights, food, rides) "
    "through an app. They expect housing to work the same way. It does "
    "not."
)

add_para(
    "Parents are also part of this. The students we interviewed who paid "
    "six months upfront were almost always doing it because their family "
    "wanted certainty. A platform that gives parents a verified booking "
    "and a contract they can read is worth a measurable premium."
)

add_heading("3.3 Our value proposition", level=2)
add_para(
    "Euronest gives a student three things that are missing from the "
    "current market:"
)
add_numbered([
    "A verified apartment. Every property is visited by a team member or "
    "a vetted local partner before it is listed. We refuse roughly a "
    "third of the apartments offered to us.",
    "A single price. Rent, utilities, internet and the chosen lifestyle "
    "services come in one monthly invoice. No separate bills, no surprise "
    "deposits, no guarantor required for the standard and premium tiers "
    "because Euronest itself acts as the guarantor.",
    "One human contact. Each subscriber gets a named coordinator who "
    "speaks at least English plus one of French, Hindi, Arabic or "
    "Mandarin. They handle paperwork in the first month and stay "
    "reachable for the rest of the stay.",
])

add_heading("3.4 What makes the offer hard to copy", level=2)
add_para(
    "Airbnb has scale but treats students like any other short-stay "
    "guest. The big French agencies (Studapart, Lokaviz, ImmoJeune) "
    "publish listings but stop at the booking. Local agencies handle "
    "paperwork but charge agency fees that match or exceed our monthly "
    "premium and they do not bundle anything beyond the apartment itself."
)
add_para(
    "What is hard to copy is the combination: a curated inventory, a "
    "guarantor function backed by our own balance sheet, and a service "
    "layer that only makes sense if you are willing to specialise in "
    "students. We are, because that is the only market we want."
)

page_break()

# =====================================================================
# 4. MARKET RESEARCH AND STRATEGY
# =====================================================================
add_heading("4. Market research and strategy", level=1)

add_heading("4.1 Market size and target cities", level=2)
add_para(
    "We are starting in five cities that together host more than 60% of "
    "international students in France. The order below also reflects the "
    "order we plan to enter them."
)

add_table(
    ["City", "Intl. students (est.)", "Avg. monthly rent (studio)", "Priority"],
    [
        ["Paris", "115,000", "EUR 950", "Year 1 launch"],
        ["Lyon", "32,000", "EUR 620", "Year 1 launch"],
        ["Lille", "21,000", "EUR 540", "Year 2"],
        ["Toulouse", "19,000", "EUR 560", "Year 2"],
        ["Marseille", "17,000", "EUR 600", "Year 3"],
    ],
)

add_para(
    "Rents in the table are city averages for an unfurnished studio. Our "
    "furnished, all-inclusive price will sit above the headline average "
    "but below what students currently pay once you add furniture, "
    "internet, gym and insurance separately."
)

add_heading("4.2 Customer segments", level=2)
add_bullets([
    "Non-EU master's students. Our primary segment. High willingness to "
    "pay for certainty, parents involved in the decision, average stay "
    "of 16 to 24 months.",
    "Erasmus students. Shorter stays of 5 to 10 months, lower price "
    "tolerance, but they bring strong word of mouth across European "
    "universities.",
    "Interns on six-month contracts. The cleanest fit for our six-month "
    "prepay product. Often arriving for a fixed start date.",
    "Young professionals new to France. Smaller volume, higher price "
    "tolerance, useful for filling apartments in the summer gap.",
])

add_heading("4.3 Competitor map", level=2)
add_table(
    ["Competitor", "Inventory", "Service depth", "Price vs Euronest"],
    [
        ["Airbnb", "Very large", "None (listings only)", "Higher per night, no contract"],
        ["Studapart", "Large", "Booking only", "Similar rent, fewer extras"],
        ["Lokaviz / ImmoJeune", "Medium", "Listings + filters", "Lower, but DIY"],
        ["Local rental agencies", "Medium", "Paperwork", "Higher once fees are added"],
        ["University residences (CROUS)", "Small per city", "Basic", "Lower price, very limited spots"],
    ],
)

add_para(
    "The honest read on this table: nobody is doing what we want to do, "
    "but pieces of it exist everywhere. Our risk is not that a competitor "
    "blocks us. It is that students do not yet trust a new brand enough "
    "to prepay six months. We address that in Section 6 under risk."
)

add_heading("4.4 Commercial strategy", level=2)
add_para(
    "Positioning. We are not the cheapest housing site and we will not "
    "pretend to be. We are the safest one for a student who has never "
    "lived in France. Every piece of marketing copy is written to that "
    "person."
)
add_para(
    "Pricing. Three tiers (basic, standard, premium) covered in detail in "
    "Section 5. A six-month prepay option carries a 5% discount and "
    "doubles as our cash-flow buffer for the first year."
)
add_para(
    "Channels. Year one is mostly digital and word of mouth: Instagram "
    "and TikTok content in English with local-language subtitles, a "
    "referral programme that pays the referring student EUR 80 per "
    "completed booking, and partnerships with five Paris schools where "
    "we already have student contacts. YouTube long-form (city guides, "
    "lease explainers) is the slow-burn channel that should mature by "
    "year two."
)
add_para(
    "Customer experience. Easy booking is the table-stakes claim. What "
    "we are actually building is the boring part: a clear lease in two "
    "languages, a deposit held in escrow, a 48-hour move-in checklist, "
    "and one named coordinator the student can reach on WhatsApp."
)

page_break()

add_heading("4.5 SWOT analysis", level=2)
add_table(
    ["Strengths", "Weaknesses"],
    [[
        "Niche focus on international students.\n"
        "Bundled offer is hard to assemble piecemeal.\n"
        "Recurring annual demand from each university intake.\n"
        "Multilingual team with first-hand experience of the problem.",
        "New brand with no trust history.\n"
        "Capital-intensive: deposits and furniture tie up cash.\n"
        "Small inventory in year one limits choice for students.\n"
        "Founder team is full-time studying for the first six months."
    ]],
)
add_table(
    ["Opportunities", "Threats"],
    [[
        "French government target of 500,000 international students by 2027.\n"
        "Expansion path through Erasmus network into Spain, Italy, Germany.\n"
        "Landlords who want stable tenants and one professional counterparty.\n"
        "AI roommate matching as a low-cost differentiator.",
        "Regulation on short-term lets and sub-letting.\n"
        "Seasonal demand: a quiet July and August reduce occupancy.\n"
        "Currency risk for students paying from outside the eurozone.\n"
        "A large player (Airbnb, Booking) copying the bundle."
    ]],
)

add_heading("4.6 PESTEL summary", level=2)
add_bullets([
    "Political. France actively recruits international students. Visa "
    "policy is the main variable to watch; a tightening would hit us "
    "directly.",
    "Economic. Inflation in rents has slowed since 2024 but utility "
    "costs remain volatile. We protect our margin by indexing the "
    "premium tier to a published utility index, reviewed every six "
    "months.",
    "Social. Students are used to subscription models for everything "
    "else in their life. Selling housing as a subscription is more "
    "familiar to them than to their parents.",
    "Technological. Booking, payments and roommate matching are off-"
    "the-shelf problems. Our differentiation is in the curation and the "
    "service, not the code.",
    "Environmental. Furnished apartments give us control over energy "
    "ratings. We will only sign with properties at DPE class D or "
    "better, which also reduces utility bills.",
    "Legal. The Alur law and the loi ELAN govern most of what we do. We "
    "have engaged a French housing lawyer for the contract templates "
    "(see Section 8).",
])

page_break()

# =====================================================================
# 5. BUSINESS MODEL
# =====================================================================
add_heading("5. Business model", level=1)

add_heading("5.1 Value created and for whom", level=2)
add_para(
    "We create value on three sides. Students get a verified place to "
    "live and a single bill. Landlords get a professional tenant, on "
    "time payments, and no vacancy management. Service partners (gyms, "
    "ISPs, meal providers, insurers) get a steady, pre-segmented stream "
    "of subscribers with negligible churn during the academic year."
)

add_heading("5.2 Customer segments", level=2)
add_para(
    "Already covered in Section 4.2. The mix we plan for in year one is "
    "roughly 60% master's students, 25% Erasmus and interns, 15% young "
    "professionals."
)

add_heading("5.3 The offer (tiers)", level=2)
add_table(
    ["Tier", "Monthly price", "What is included", "Best for"],
    [
        ["Basic", "EUR 690 / month",
         "Furnished apartment, fibre internet, utilities, contents insurance.",
         "Short stays, budget-sensitive Erasmus students."],
        ["Standard", "EUR 890 / month",
         "Everything in Basic + gym pass, Netflix or Amazon Prime, "
         "Euronest acts as guarantor.",
         "Master's students for a full academic year."],
        ["Premium", "EUR 1,190 / month",
         "Everything in Standard + meal plan (15 meals/week), visa & "
         "paperwork support, swimming pool / sports turf access, one "
         "annual return-home flight credit (capped at EUR 450).",
         "Students whose families want the full setup."],
    ],
)
add_para(
    "Prices above are Paris pricing. Lyon and the other cities are 15 to "
    "25% lower because the underlying rent is lower. The six-month "
    "prepay option carries a 5% discount on top of these prices."
)

add_heading("5.4 Revenue streams", level=2)
add_bullets([
    "Subscription revenue (90% of total). Monthly or six-month prepay.",
    "Landlord commission (6%). On apartments we list but do not lease "
    "ourselves, we keep a commission on the first month of rent.",
    "Service partner share. Gyms, ISPs and meal providers pay us a "
    "fixed margin on each subscriber we route to them.",
    "Gift vouchers and family top-ups. Parents buy credit that students "
    "spend inside the platform.",
    "Optional add-ons. Airport pickup, French SIM card, additional "
    "language coaching. Small in revenue, high in customer satisfaction.",
])

add_heading("5.5 Cost structure", level=2)
add_table(
    ["Cost item", "Type", "Year 1 estimate (EUR)"],
    [
        ["Apartment rent paid to landlords", "Variable", "240,000"],
        ["Furniture and setup (recoverable)", "One-off", "55,000"],
        ["Platform and mobile app development", "One-off", "35,000"],
        ["Staff salaries (4 FTE by month 6)", "Fixed", "110,000"],
        ["Cleaning and maintenance", "Variable", "18,000"],
        ["Marketing and advertising", "Variable", "42,000"],
        ["Office rent and utilities", "Fixed", "9,000"],
        ["Insurance, legal, accounting", "Fixed", "14,000"],
        ["Bank fees and payment processing", "Variable", "6,000"],
        ["Total year 1", "", "529,000"],
    ],
)
add_para(
    "The largest line is rent we pay to landlords, which is matched by "
    "the rent we collect from subscribers. The number that actually "
    "needs to come from investment is the gap between cumulative costs "
    "and cumulative revenue during the first nine months. We size that "
    "gap at around EUR 180,000 in Section 7."
)

add_heading("5.6 Key activities", level=2)
add_bullets([
    "Sourcing and verifying apartments.",
    "Furnishing and setting up each unit before the academic year.",
    "Onboarding students, including paperwork and arrival logistics.",
    "Running the website, app and payment infrastructure.",
    "Managing service partners and renewing volume contracts.",
])

add_heading("5.7 Key resources", level=2)
add_bullets([
    "Verified apartment inventory (the moat that compounds over years).",
    "Brand and reputation among university communities.",
    "The platform itself: bookings, payments, escrow, roommate matching.",
    "Working capital to pre-pay landlord deposits.",
    "The founding team's network across student associations.",
])

add_heading("5.8 Key partners", level=2)
add_bullets([
    "Landlords and small property managers in our five target cities.",
    "ISPs (Orange, Free, SFR) for the internet bundle.",
    "Gym chains (Basic-Fit, Fitness Park, On Air Fitness) for the "
    "lifestyle bundle.",
    "Insurers for contents and civil liability cover.",
    "Banks for student-friendly account opening (Boursorama, Revolut, "
    "BNP Paribas).",
    "Universities and student associations for distribution.",
])

add_heading("5.9 Operational workflow", level=2)
add_numbered([
    "Sourcing. The operations lead and one apartment scout visit each "
    "candidate unit, check it against a 38-point checklist, and "
    "negotiate a 12 to 24 month head-lease with the landlord.",
    "Setup. The apartment is furnished from a standard catalogue (bed, "
    "desk, chair, kitchenware) and fitted with fibre internet within "
    "two weeks of signing.",
    "Listing. Photos, floor plan and a short walk-through video go on "
    "the website and the app. Pricing follows the tier table in 5.3.",
    "Booking. The student selects a tier, signs the bilingual lease, "
    "and pays either the first month plus deposit or the six-month "
    "upfront amount.",
    "Move-in. The coordinator meets the student at the apartment, "
    "hands over keys, walks through the inventory, and registers the "
    "student with the local utilities where required.",
    "Support. Coordinator stays on WhatsApp for the duration of the "
    "stay. Maintenance issues are routed to vetted local handymen with "
    "a 48-hour response target.",
])

page_break()

# =====================================================================
# 6. HUMAN RESOURCES, RISKS AND MITIGATION
# =====================================================================
add_heading("6. Human resources, risks and mitigation", level=1)

add_heading("6.1 The founding team", level=2)
add_para(
    "We are seven MBA students at ISTEC. None of us has run a housing "
    "business before, but between us we have worked in software, "
    "finance, marketing and hospitality, and we have all been on the "
    "wrong side of the problem we are now trying to solve."
)

add_table(
    ["Member", "Role in Euronest", "Background"],
    [
        ["Mansoor Ali", "CEO, partnerships",
         "Hospitality operations; led a 40-room boutique hotel before MBA."],
        ["Fidha Fathima Panankavil", "COO, apartment operations",
         "Real estate analyst; handles inventory, setup, maintenance."],
        ["Leo Joy", "CFO",
         "Audit and corporate finance; runs the financial model and fundraising."],
        ["Nithin Stanley George", "CTO, platform",
         "Software engineer; owns the website, app and payments."],
        ["Minhaj Eshwaramangalam", "Head of marketing",
         "Digital marketing; runs paid social, content and the referral programme."],
        ["Gopinath Dasan", "Head of student support",
         "Multilingual customer service background; coordinator team lead."],
        ["Pallampettivelayudhan Nithin", "Head of legal and compliance",
         "Law graduate; manages lease templates, GDPR, regulatory tracking."],
    ],
)

add_heading("6.2 Hiring plan", level=2)
add_para(
    "We stay at seven founders through the pilot. Once we cross 60 paid "
    "subscribers we add three roles: two student coordinators (one for "
    "Paris, one for Lyon) and one apartment scout. By the end of year "
    "two we expect twelve people on payroll. Salaries are modelled "
    "against the French market median for each role with a 10% discount "
    "compensated by equity for the first four hires."
)

add_heading("6.3 Culture and how we work", level=2)
add_para(
    "We are remote-first with a shared office at ISTEC for two days a "
    "week. Operations meet daily for 15 minutes. Everyone, including "
    "the engineers, spends at least one day a month visiting apartments "
    "or meeting students. This is a deliberate rule. The moment a "
    "Euronest employee stops talking to the customer is the moment we "
    "lose what made the product different in the first place."
)

add_heading("6.4 Risks", level=2)
add_para(
    "Listing them honestly. No risk register is complete, but these are "
    "the ones we think about most."
)

add_table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Low trust in a new brand slows year-1 sign-ups",
         "High", "High",
         "Money-back guarantee on the first month. Verified reviews. "
         "First-cohort discount of 10% for the September 2026 intake."],
        ["Landlord pulls a property mid-lease",
         "Medium", "Medium",
         "12-month head-leases with break clauses on our side only. "
         "Backup inventory of 10% spare units per city."],
        ["Regulatory change on furnished rentals",
         "Low", "High",
         "Quarterly review with French housing lawyer. Lobby through "
         "FNAIM membership."],
        ["Cash gap before break-even",
         "Medium", "High",
         "Six-month prepay option creates working capital. Reserve of "
         "EUR 40k held back from initial raise."],
        ["Key founder leaves before year-2",
         "Medium", "Medium",
         "Vesting on four-year schedule. Two-deep coverage for every "
         "function."],
        ["Apartment damage or non-payment by subscriber",
         "Medium", "Low to medium",
         "Two-month deposit held in escrow. Civil liability insurance "
         "bundled into the standard tier."],
        ["Currency volatility for non-EU parents paying in advance",
         "Low", "Medium",
         "Lock-in exchange rate via Wise Business at the time of booking."],
    ],
)

add_heading("6.5 What would make us shut this down", level=2)
add_para(
    "If, by the end of month 18, we have fewer than 120 active "
    "subscribers and a customer acquisition cost above EUR 280, we "
    "stop. The platform either compounds on word of mouth or it does "
    "not, and pouring more marketing into a product students do not "
    "recommend would be the wrong call."
)

page_break()

# =====================================================================
# 7. FINANCIAL FORECASTS
# =====================================================================
add_heading("7. Financial forecasts", level=1)

add_heading("7.1 Initial investment", level=2)
add_table(
    ["Use of funds", "Amount (EUR)"],
    [
        ["Platform and mobile app development", "35,000"],
        ["Furniture, deposits and setup for first 25 apartments", "55,000"],
        ["Marketing for September 2026 intake", "25,000"],
        ["Six months of fixed operating costs", "45,000"],
        ["Working capital reserve", "20,000"],
        ["Total", "180,000"],
    ],
)
add_para(
    "We are raising the EUR 180,000 as a mix of founder contribution "
    "(EUR 40k), a Bpifrance student entrepreneur loan (EUR 60k applied "
    "for in March 2026) and an equity round from one French business "
    "angel network for the balance (EUR 80k for 12% equity)."
)

add_heading("7.2 Three-year projection", level=2)
add_table(
    ["", "Year 1", "Year 2", "Year 3"],
    [
        ["Subscribers (avg.)", "55", "180", "320"],
        ["Mix (basic / std / prem)", "35/45/20%", "25/55/20%", "20/55/25%"],
        ["Revenue (EUR)", "540,000", "1,820,000", "3,360,000"],
        ["Cost of services (rent + bundles)", "385,000", "1,210,000", "2,180,000"],
        ["Gross margin", "155,000", "610,000", "1,180,000"],
        ["Gross margin %", "29%", "34%", "35%"],
        ["Operating costs (team + ops)", "190,000", "395,000", "640,000"],
        ["EBITDA", "-35,000", "215,000", "540,000"],
        ["EBITDA %", "-6%", "12%", "16%"],
    ],
)

add_para(
    "Year 1 is intentionally loss-making. We are buying the first cohort "
    "with the September intake discount and we are paying a full team "
    "from month six. The model turns positive between month 13 and "
    "month 15 depending on how the June-August occupancy lands."
)

add_heading("7.3 Break-even analysis", level=2)
add_para(
    "Fixed monthly costs after the team is fully hired sit at around "
    "EUR 18,500 (salaries, office, software, accounting, marketing "
    "baseline). Average contribution per subscriber is EUR 480 per "
    "month at the year-1 tier mix. That puts break-even at 39 "
    "subscribers, which matches the 35 to 45 range we worked through in "
    "the financial overview meeting."
)

add_heading("7.4 Cash flow", level=2)
add_para(
    "The six-month prepay option is doing real work on the cash flow "
    "side. Even if only a quarter of standard and premium subscribers "
    "take it, that pulls forward roughly EUR 95,000 of cash in the "
    "first quarter of operation. That is the same money that would "
    "otherwise need to come from a larger raise."
)

add_para(
    "We model a worst-case scenario where uptake is 40% slower than "
    "planned. In that case the company still survives year one but "
    "needs an additional EUR 60,000 bridge in month 9. We have "
    "pre-agreed this bridge in principle with the angel investor."
)

add_heading("7.5 Sensitivities", level=2)
add_bullets([
    "If average rent paid to landlords rises by 8%, EBITDA in year 2 "
    "drops from EUR 215,000 to around EUR 130,000. Mitigation: longer "
    "head-leases with fixed rent.",
    "If gym partners pull their volume discount, the standard tier "
    "loses EUR 22 per subscriber per month. We have a backup partner "
    "agreement with On Air Fitness signed in February 2026.",
    "If churn between semesters rises from 12% to 25%, year-3 revenue "
    "drops by roughly EUR 480,000. Mitigation: annual prepay incentive "
    "and end-of-semester roommate matching for the next cohort.",
])

page_break()

# =====================================================================
# 8. LEGAL AND ORGANISATIONAL STRUCTURE
# =====================================================================
add_heading("8. Legal and organisational structure", level=1)

add_heading("8.1 Legal form", level=2)
add_para(
    "Euronest will be registered as a SAS (Societe par Actions "
    "Simplifiee) under French law. We picked SAS over SARL for three "
    "reasons: flexibility in shareholder agreements, easier issuance of "
    "founder shares and BSPCE (employee equity), and the option to "
    "bring in a venture investor later without restructuring."
)

add_heading("8.2 Shareholding at incorporation", level=2)
add_table(
    ["Shareholder", "Shares", "Percentage"],
    [
        ["Mansoor Ali", "13,000", "13.0%"],
        ["Fidha Fathima Panankavil", "13,000", "13.0%"],
        ["Leo Joy", "13,000", "13.0%"],
        ["Nithin Stanley George", "13,000", "13.0%"],
        ["Minhaj Eshwaramangalam", "13,000", "13.0%"],
        ["Gopinath Dasan", "13,000", "13.0%"],
        ["Pallampettivelayudhan Nithin", "10,000", "10.0%"],
        ["Angel investor (post-seed)", "12,000", "12.0%"],
        ["Total", "100,000", "100%"],
    ],
)
add_para(
    "All founder shares vest over four years with a one-year cliff. The "
    "angel investor's stake is post-seed and reflects the EUR 80,000 "
    "investment described in Section 7."
)

add_heading("8.3 Governance", level=2)
add_para(
    "A president (Mansoor Ali, by founder vote) and a directeur general "
    "(Leo Joy, also CFO) are the two statutory officers. Major "
    "decisions (raising more equity, opening a new country, executive "
    "hires above EUR 70,000) require a two-thirds founder vote. The "
    "angel investor has an information right but no veto."
)

add_heading("8.4 Internal organisation", level=2)
add_bullets([
    "Operations team (apartment sourcing, setup, maintenance) reports "
    "to the COO.",
    "Student support team (coordinators in each city) reports to the "
    "head of student support.",
    "Platform team (engineers, designer) reports to the CTO.",
    "Marketing and growth report to the head of marketing.",
    "Finance, legal and HR report to the CFO and the head of legal "
    "respectively, both reporting to the CEO.",
])

add_heading("8.5 Regulatory and tax", level=2)
add_para(
    "Euronest is a furnished rental operator subject to the loi ALUR "
    "and, where applicable, the loi ELAN. We register under the "
    "regime LMNP (Loueur en Meuble Non Professionnel) for year one and "
    "expect to move to LMP from year two onwards. VAT applies on the "
    "service portion of the bundle (gym, streaming, food) but not on "
    "the rent itself, in line with French tax practice."
)

add_heading("8.6 Data protection", level=2)
add_para(
    "GDPR applies to everything we store about students. We have "
    "appointed Pallampettivelayudhan Nithin as data protection officer. "
    "Student data is hosted on a French cloud provider (OVH) with "
    "encryption at rest. We do not share data with service partners "
    "beyond the strict minimum needed to provision the bundled service."
)

page_break()

# =====================================================================
# 9. ROADMAP
# =====================================================================
add_heading("9. Roadmap", level=1)

add_heading("9.1 Milestones", level=2)
add_table(
    ["Quarter", "Milestone", "Owner"],
    [
        ["Q2 2026", "Company incorporated as SAS. Initial raise closed.",
         "CEO / CFO"],
        ["Q2 2026", "MVP of the booking website live in English and French.",
         "CTO"],
        ["Q3 2026", "First 25 apartments under head-lease in Paris and Lyon.",
         "COO"],
        ["Q3 2026", "September 2026 intake: 20 paid subscribers onboarded.",
         "Marketing / Support"],
        ["Q4 2026", "Mobile app (iOS and Android) released.", "CTO"],
        ["Q1 2027", "55 active subscribers. First three-month financial review.",
         "CFO"],
        ["Q2 2027", "Lille and Toulouse pilot. 100 active subscribers.",
         "COO / Marketing"],
        ["Q3 2027", "AI roommate matching module live.", "CTO"],
        ["Q4 2027", "180 active subscribers. EBITDA positive month achieved.",
         "Whole team"],
        ["Q2 2028", "Marseille launched. Spain market study begins.",
         "CEO"],
        ["Q4 2028", "320 active subscribers. Series A discussions opened.",
         "CEO / CFO"],
    ],
)

add_heading("9.2 What we will know by month 18", level=2)
add_para(
    "By the end of 2027 we will know three things: whether the "
    "subscription bundle holds together at scale (margin discipline), "
    "whether students stay or churn between academic years (retention), "
    "and whether the brand carries across cities or needs to be "
    "rebuilt locally (marketing efficiency). These three answers decide "
    "whether year three is an expansion year or a consolidation year."
)

add_heading("9.3 Beyond France", level=2)
add_para(
    "The natural next markets are Spain (Madrid, Barcelona) and Italy "
    "(Milan), both of which share the same Erasmus and non-EU student "
    "patterns. Germany is larger but more regulated. We will not "
    "commit to a country outside France before we are comfortably "
    "EBITDA-positive at home. The plan is to grow boring before we "
    "grow ambitious."
)

page_break()

# =====================================================================
# 10. APPENDIX
# =====================================================================
add_heading("10. Appendix", level=1)

add_heading("A. Glossary", level=2)
add_bullets([
    "ALUR / ELAN: French housing laws governing rental contracts and "
    "tenant protections.",
    "BSPCE: bons de souscription de parts de createur d'entreprise. "
    "The French equivalent of employee stock options.",
    "DPE: Diagnostic de Performance Energetique. Energy rating for a "
    "French property.",
    "LMNP / LMP: tax regimes for non-professional and professional "
    "furnished rental operators.",
    "SAS: Societe par Actions Simplifiee. A flexible French corporate "
    "form often used by startups.",
])

add_heading("B. Source meeting minutes", level=2)
add_bullets([
    "Session 3 (Problem Diagnosis and Business Strategy): problem "
    "analysis, why now, market research, commercial strategy, SWOT and "
    "PESTEL analysis.",
    "Session on Business Model and Financial Planning: value "
    "proposition, revenue model, cost structure, key partners, "
    "operational workflow, financial overview.",
    "New Venture Development meeting (10 April 2026, ISTEC): problem "
    "identification, business idea, USP, tier-based model, AI roommate "
    "matching.",
])

add_heading("C. References to the Business Creation Process", level=2)
add_para(
    "The structure of this document follows the nine-step business "
    "creation process used in the MBA English programme: cover page "
    "and contents, executive summary, problem analysis and value "
    "proposition, market research and strategy, business model, human "
    "resources and risks, financial forecasts, legal and organisational "
    "structure, roadmap, and appendix. Each section in the body of "
    "this plan maps one-to-one to that outline."
)

add_heading("D. Contact", level=2)
add_para(
    "Euronest founding team\nISTEC, Paris\nfounders@euronest.eu"
)

# =====================================================================
# Save
# =====================================================================
out = "/home/user/first/Euronest_Business_Plan.docx"
doc.save(out)
print(f"Saved: {out}")
